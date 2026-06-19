from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Function to set borders for a table
def set_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            cell_pr = cell._element.get_or_add_tcPr()
            if cell_pr.find(qn('w:tcBorders')) is None:
                cell_borders = OxmlElement('w:tcBorders')
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')  # Border style
                    border.set(qn('w:sz'), '4')  # Border size
                    cell_borders.append(border)
                cell_pr.append(cell_borders)

# Create a new Document
doc = Document()
doc.add_heading('Database Schema', 0)

# Add Users Table
doc.add_heading('Users Table', level=1)

# Define Users Table Schema
users_table_data = [
    ['Column Name', 'Data Type', 'Constraints', 'Description'],
    ['id', 'INTEGER', 'PRIMARY KEY, AUTOINCREMENT', 'Unique identifier for each user'],
    ['username', 'VARCHAR(50)', 'UNIQUE, NOT NULL', 'User’s login username'],
    ['password', 'VARCHAR(255)', 'NOT NULL', 'User’s hashed password']
]

# Add Users Table to Document
users_table = doc.add_table(rows=1, cols=4)
hdr_cells = users_table.rows[0].cells
for i, header in enumerate(users_table_data[0]):
    hdr_cells[i].text = header

for row_data in users_table_data[1:]:
    row_cells = users_table.add_row().cells
    for i, data in enumerate(row_data):
        row_cells[i].text = data

# Set borders for Users Table
set_table_borders(users_table)

# Add Predictions Table
doc.add_heading('Predictions Table', level=1)

# Define Predictions Table Schema
predictions_table_data = [
    ['Column Name', 'Data Type', 'Constraints', 'Description'],
    ['id', 'INTEGER', 'PRIMARY KEY, AUTOINCREMENT', 'Unique identifier for each prediction'],
    ['user_id', 'INTEGER', 'FOREIGN KEY (references `users.id`), NOT NULL', 'Identifier for the user making the prediction'],
    ['spx', 'FLOAT', 'NOT NULL', 'Value for SPX'],
    ['uso', 'FLOAT', 'NOT NULL', 'Value for USO'],
    ['slv', 'FLOAT', 'NOT NULL', 'Value for SLV'],
    ['eur_usd', 'FLOAT', 'NOT NULL', 'Value for EUR/USD'],
    ['prediction', 'FLOAT', 'NOT NULL', 'Predicted gold price'],
    ['timestamp', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP', 'Time when the prediction was made']
]

# Add Predictions Table to Document
predictions_table = doc.add_table(rows=1, cols=4)
hdr_cells = predictions_table.rows[0].cells
for i, header in enumerate(predictions_table_data[0]):
    hdr_cells[i].text = header

for row_data in predictions_table_data[1:]:
    row_cells = predictions_table.add_row().cells
    for i, data in enumerate(row_data):
        row_cells[i].text = data

# Set borders for Predictions Table
set_table_borders(predictions_table)

# Save the Document
doc.save('Database_Schema_Bordered.docx')

print("Document created and saved as 'Database_Schema_Bordered.docx'")
